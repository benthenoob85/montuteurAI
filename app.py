import streamlit as st
import google.generativeai as genai
from groq import Groq
from pypdf import PdfReader
import pandas as pd
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import io
import matplotlib.pyplot as plt
import re
import tempfile
import os

# --- LIBRAIRIES RAG (IA Documentaire) ---
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferWindowMemory

# --- 1. CONFIGURATION INITIALE ---
st.set_page_config(page_title="Tuteur Financier RAG", layout="wide", page_icon="🎓")

# Styles pour les badges et l'interface
st.markdown("""
<style>
    .stChatMessage {background-color: #f0f2f6; border-radius: 10px; padding: 10px; margin-bottom: 10px;}
    .badge {padding: 3px 8px; border-radius: 4px; font-size: 0.8em; font-weight: bold; color: white;}
    .badge-flash {background-color: #28a745;}
    .badge-pro {background-color: #007bff;}
    .badge-groq {background-color: #dc3545;}
    /* Ajustement des boutons pour l'export */
    .stDownloadButton > button {height: auto; padding: 5px; font-size: 0.8em; width: 100%;}
    .stButton button {width: 100%;}
</style>
""", unsafe_allow_html=True)

# --- 2. GESTION DES SECRETS ET CONNEXIONS ---
valid_google_models = []
groq_client = None

# Connexion Google (avec détection des modèles)
try:
    if "GOOGLE_API_KEY" in st.secrets:
        genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
        try:
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    valid_google_models.append(m.name)
        except:
            valid_google_models = ["models/gemini-flash-latest", "models/gemini-pro"]
except: pass

# Connexion Groq
if "GROQ_API_KEY" in st.secrets:
    try: groq_client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    except: pass

# --- 3. GESTION DU RAG ET DES SESSIONS ---
if 'chat_sessions' not in st.session_state:
    st.session_state['chat_sessions'] = {} # Archive des conversations
if 'current_session_id' not in st.session_state:
    st.session_state['current_session_id'] = "Session 1"
if 'vector_store' not in st.session_state:
    st.session_state['vector_store'] = None # Base de vecteurs pour le RAG
if 'context_text' not in st.session_state:
    st.session_state['context_text'] = "" # Texte brut des documents
if 'memory' not in st.session_state:
    # Mémoire glissante de 5 échanges (limitée en tokens)
    st.session_state['memory'] = ConversationBufferWindowMemory(k=5, memory_key="chat_history", return_messages=True)

# --- CACHING DES EMBEDDINGS (Performance) ---
@st.cache_resource(show_spinner=False)
def get_vector_store(uploaded_file_names, context_text):
    """Crée le vector store RAG et le met en cache"""
    if not context_text: return None

    # 1. Découpage du texte (Chunking)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(context_text)

    # 2. Création des Embeddings (vecteurs) avec Google
    st.toast("Création des embeddings avec Google...", icon="🧠")
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", client=genai)
    
    # 3. Création du Vector Store (FAISS est léger et rapide)
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    return vector_store

# --- ARCHIVAGE DES SESSIONS ---
def start_new_session():
    """Archive la session actuelle et en commence une nouvelle"""
    current_messages = st.session_state.get('messages', [])
    if current_messages:
        # Stocker la session actuelle dans l'archive
        session_name = st.session_state['current_session_id']
        st.session_state['chat_sessions'][session_name] = current_messages
    
    # Réinitialisation
    session_id = f"Session {len(st.session_state['chat_sessions']) + 1}"
    st.session_state['current_session_id'] = session_id
    st.session_state['messages'] = []
    st.session_state['memory'].clear()
    st.toast(f"Nouvelle session '{session_id}' démarrée!", icon="💬")

def load_session(session_id):
    """Charge une session archivée"""
    st.session_state['current_session_id'] = session_id
    st.session_state['messages'] = st.session_state['chat_sessions'][session_id]
    st.session_state['memory'].clear() # On vide la mémoire de travail
    st.toast(f"Session '{session_id}' chargée!", icon="📂")

# --- LECTURE ET PRÉ-TRAITEMENT DES FICHIERS ---
def load_and_process_files(uploaded_files):
    raw_text = ""
    for uploaded_file in uploaded_files:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            if file_type == 'pdf':
                loader = PyPDFLoader(tmp_file_path)
            elif file_type == 'docx':
                loader = Docx2txtLoader(tmp_file_path)
            elif file_type in ['txt']:
                loader = TextLoader(tmp_file_path)
            elif file_type in ['xlsx', 'xls']:
                xls = pd.ExcelFile(tmp_file_path)
                excel_text = ""
                for sheet in xls.sheet_names:
                    excel_text += f"\n--- Excel: {sheet} ---\n" + pd.read_excel(xls, sheet_name=sheet).to_string()
                raw_text += excel_text
                continue 
            elif file_type in ['png', 'jpg', 'jpeg']:
                image = Image.open(uploaded_file)
                vision_model_name = get_google_model_name("flash")
                vision_model = genai.GenerativeModel(vision_model_name)
                response = vision_model.generate_content(["Transcris tout le texte :", image])
                raw_text += f"\n--- Image ---\n{response.text}"
                continue
            else:
                st.warning(f"Fichier ignoré (format non supporté) : .{file_type}")
                continue

            docs = loader.load()
            for doc in docs:
                raw_text += doc.page_content
        
        except Exception as e:
            st.error(f"Erreur de lecture RAG pour {uploaded_file.name}: {e}")
        finally:
            if 'tmp_file_path' in locals() and os.path.exists(tmp_file_path):
                 os.remove(tmp_file_path)
                 
    if raw_text:
        st.session_state['context_text'] = raw_text
        st.session_state['vector_store'] = get_vector_store(
            [f.name for f in uploaded_files], 
            raw_text
        )
        st.success("Documents chargés et indexés pour le RAG!")

# --- 4. FONCTION PRINCIPALE D'INTERROGATION (CASCADE RAG) ---
def get_google_model_name(type="flash"):
    """Trouve le meilleur nom de modèle valide"""
    clean_models = [m.replace('models/', '') for m in valid_google_models]
    
    if type == "pro":
        candidates = [m for m in clean_models if 'pro' in m and '2.5' in m]
        if not candidates: candidates = [m for m in clean_models if 'pro' in m and 'latest' in m]
        return candidates[0] if candidates else 'gemini-pro'
    else: 
        candidates = [m for m in clean_models if 'flash' in m and 'latest' in m]
        if not candidates: candidates = [m for m in clean_models if 'flash' in m]
        return candidates[0] if candidates else 'gemini-flash-latest'

def ask_groq(prompt, system_instruction):
    if not groq_client: raise Exception("Pas de clé Groq")
    chat = groq_client.chat.completions.create(
        messages=[{"role": "system", "content": system_instruction}, {"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile", temperature=0
    )
    return chat.choices[0].message.content

def ask_smart_ai(prompt):
    """Gère la cascade AI et le RAG"""
    system_instruction = (
        "Tu es un expert pédagogique Finance. Règle stricte pour la compatibilité PDF/Word: "
        "Toutes les variables, symboles, lettres grecques (alpha, sigma, etc.) ou expressions mathématiques "
        "DOIVENT être entourées de $$...$$ sur une ligne séparée (même les symboles simples) pour garantir le bon export. "
        "Exemple: 'La volatilité $$ \\sigma $$ est calculée par...' devient 'La volatilité $$ \\sigma $$ est calculée...'"
    )
    full_prompt = prompt

    # --- ÉTAPE RAG ---
    if st.session_state['vector_store']:
        retriever = st.session_state['vector_store'].as_retriever()
        docs = retriever.get_relevant_documents(prompt)
        context = "\n\n--- CONTEXTE RÉCUPÉRÉ ---\n" + "\n".join([doc.page_content for doc in docs])
        full_prompt = f"Question : {prompt}\n\nContexte : {context}"
        is_complex = True
    else:
        complexity_keywords = ["analyse", "synthèse", "résous", "calcul", "tableau", "excel", "bilan", "ratio", "formule", "développe", "argumente", "comparer", "pourquoi"]
        is_complex = any(k in prompt.lower() for k in complexity_keywords) or (len(prompt.split()) > 15)

    # --- CASCADE ---
    if not is_complex:
        model_name = get_google_model_name("flash")
        cascade = [(model_name, "⚡ Gemini Flash"), ('llama-3.3-70b-versatile', "🦙 Groq (Secours Flash)")]
    else:
        pro_model = get_google_model_name("pro")
        flash_model = get_google_model_name("flash")
        cascade = [
            (pro_model, "🧠 Gemini Pro"),
            ('llama-3.3-70b-versatile', "🦙 Groq (Relais Pro)"),
            (flash_model, "⚡ Gemini Flash (Secours Ultime)")
        ]
    
    last_error = ""
    for model_name, label in cascade:
        try:
            if "llama" in model_name:
                resp = ask_groq(full_prompt, system_instruction)
            else:
                llm = ChatGoogleGenerativeAI(model=model_name, temperature=0, client=genai)
                conversation = RetrievalQA.from_chain_type(
                    llm=llm,
                    retriever=retriever if st.session_state['vector_store'] else None,
                    memory=st.session_state['memory'],
                    chain_type="stuff",
                    verbose=False,
                    chain_type_kwargs={"prompt": PromptTemplate(template=f"{system_instruction}\n\n{{context}}\n\n{{question}}", input_variables=["context", "question"])}
                )
                response_obj = conversation({"query": prompt})
                resp = response_obj["result"]
            return resp, label
        except Exception as e:
            last_error = e
            continue
            
    return f"Panne totale. (Dernière erreur: {last_error})", "❌ Erreur"

# --- 5. EXPORTS WORD ET PDF ---
def latex_to_image(latex_str):
    try:
        fig, ax = plt.subplots(figsize=(6, 1.5))
        clean_latex = latex_str.replace(r'\ ', ' ') 
        ax.text(0.5, 0.5, f"${clean_latex}$", size=20, ha='center', va='center')
        ax.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=300, transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf
    except: return None

def clean_text_for_word(text):
    text = text.replace('$', '')
    replacements = {
        r'\sigma': 'σ', r'\Sigma': 'Σ', r'\mu': 'μ', r'\beta': 'β', r'\alpha': 'α',
        r'\gamma': 'γ', r'\lambda': 'λ', r'\sum': '∑', r'\approx': '≈', r'\times': '×',
        r'\le': '≤', r'\ge': '≥', r'\infty': '∞', r'_i': 'ᵢ', r'_t': 'ₜ', r'_0': '₀', r'^2': '²', r'\%': '%'
    }
    for latex, char in replacements.items(): text = text.replace(latex, char)
    return text

def create_word_docx(text_content, title="Document IA"):
    doc = Document()
    doc.add_heading(title, 0)
    parts = re.split(r'(\$\$.*?\$\$)', text_content, flags=re.DOTALL)
    for part in parts:
        if part.startswith('$$') and part.endswith('$$'):
            image_buffer = latex_to_image(part.replace('$$', '').strip())
            if image_buffer: doc.add_picture(image_buffer, width=Inches(2.5))
            else: doc.add_paragraph(part)
        else:
            if part.strip(): doc.add_paragraph(clean_text_for_word(part))
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_pdf(text_content, title="Document IA"):
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 15)
            self.cell(0, 10, title, 0, 1, 'C')
            self.ln(5)
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    parts = re.split(r'(\$\$.*?\$\$)', text_content, flags=re.DOTALL)
    for part in parts:
        if part.startswith('$$') and part.endswith('$$'):
            image_buffer = latex_to_image(part.replace('$$', '').strip())
            if image_buffer:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                    temp_img.write(image_buffer.getvalue())
                    temp_img_path = temp_img.name
                try: pdf.image(temp_img_path, w=100, x=55)
                except: pass
                try: os.remove(temp_img_path)
                except: pass
        else:
            clean_txt = clean_text_for_word(part)
            # Encodage latin-1 pour FPDF (gère les accents français basiques)
            clean_txt = clean_txt.encode('latin-1', 'replace').decode('latin-1')
            if clean_txt.strip():
                pdf.multi_cell(0, 7, txt=clean_txt)
                pdf.ln(2)
    return pdf.output(dest='S').encode('latin-1')

# --- 6. INTERFACE STREAMLIT ---
with st.sidebar:
    st.header(f"Session : {st.session_state['current_session_id']}")
    st.button("🆕 Nouvelle Conversation", on_click=start_new_session, use_container_width=True)
    
    st.markdown("### 💾 Fichiers & RAG")
    uploaded_files = st.file_uploader("Documents", accept_multiple_files=True)
    if uploaded_files and st.button("🔄 Charger & Indexer"):
        with st.spinner("Indexation des documents..."):
            load_and_process_files(uploaded_files)

    if st.session_state['vector_store']: st.success("✅ RAG ACTIF")
    else: st.warning("⚠️ RAG INACTIF")

    st.markdown("### 📂 Historique")
    if st.session_state['chat_sessions']:
        for session_name in sorted(st.session_state['chat_sessions'].keys(), reverse=True):
            if session_name != st.session_state['current_session_id']:
                st.button(f"Charger {session_name}", key=session_name, on_click=load_session, args=(session_name,), use_container_width=True)
    
    st.markdown("---")
    if groq_client: st.success("✅ Groq Connecté")
    else: st.warning("❌ Groq Absent")

st.subheader(f"🎓 Tuteur Financier (RAG Actif : {'Oui' if st.session_state['vector_store'] else 'Non'})")
tab1, tab2, tab3 = st.tabs(["💬 Chat", "📝 Synthèses", "🧠 Quiz"])

with tab1:
    if "messages" not in st.session_state: st.session_state['messages'] = []
    
    for i, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant":
                used_model = msg.get("model_label", "IA")
                if "Groq" in used_model: badge = "badge-groq"
                elif "Pro" in used_model: badge = "badge-pro"
                else: badge = "badge-flash"
                st.markdown(f'<span class="badge {badge}">{used_model}</span>', unsafe_allow_html=True)
                
                c1, c2 = st.columns(2)
                with c1:
                    docx = create_word_docx(msg["content"], title=f"Réponse {i}")
                    st.download_button("💾 Word", docx.getvalue(), f"note_{i}.docx", key=f"d{i}w")
                with c2:
                    try:
                        pdf_data = create_pdf(msg["content"], title=f"Réponse {i}")
                        st.download_button("📄 PDF", pdf_data, f"note_{i}.pdf", key=f"d{i}p")
                    except: pass

    if user := st.chat_input("Votre question..."):
        st.session_state.messages.append({"role": "user", "content": user})
        with st.chat_message("user"): st.markdown(user)
        with st.chat_message("assistant"):
            with st.spinner("Analyse..."):
                resp, model_label = ask_smart_ai(user)
                st.markdown(resp)
                st.session_state.messages.append({"role": "assistant", "content": resp, "model_label": model_label})
                
                if "Groq" in model_label: badge = "badge-groq"
                elif "Pro" in model_label: badge = "badge-pro"
                else: badge = "badge-flash"
                st.markdown(f'<span class="badge {badge}">{model_label}</span>', unsafe_allow_html=True)
                
                c1, c2 = st.columns(2)
                with c1:
                    docx = create_word_docx(resp, title="Réponse")
                    st.download_button("💾 Word", docx.getvalue(), "reponse.docx", key="new_w")
                with c2:
                    try:
                        pdf_data = create_pdf(resp, title="Réponse")
                        st.download_button("📄 PDF", pdf_data, "reponse.pdf", key="new_p")
                    except: pass

with tab2:
    if st.button("Générer Synthèse"):
        if st.session_state['vector_store']:
            with st.spinner("Rédaction..."):
                resp, label = ask_smart_ai("Fais une synthèse structurée des documents.")
                st.markdown(resp)
                c1, c2 = st.columns(2)
                with c1: st.download_button("💾 Word", create_word_docx(resp).getvalue(), "synthese.docx")
                with c2: st.download_button("📄 PDF", create_pdf(resp), "synthese.pdf")
        else: st.error("Chargez des documents d'abord.")

with tab3:
    if st.button("Lancer Quiz"):
        if st.session_state['vector_store']:
            with st.spinner("Création Quiz..."):
                resp, label = ask_smart_ai("Crée 3 QCM basés sur les documents.")
                st.markdown(resp)
                c1, c2 = st.columns(2)
                with c1: st.download_button("💾 Word", create_word_docx(resp).getvalue(), "quiz.docx")
                with c2: st.download_button("📄 PDF", create_pdf(resp), "quiz.pdf")
        else: st.error("Chargez des documents d'abord.")
